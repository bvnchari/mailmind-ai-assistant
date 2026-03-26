import streamlit as st
from engine import mailmind_app, test_smtp_connection, send_real_email
from docx import Document 
from fpdf import FPDF
import pandas as pd
import io
import uuid
import os

st.set_page_config(page_title="MailMind AI", layout="wide")

# --- SESSION STATE INITIALIZATION ---
if "total_spent" not in st.session_state:
    st.session_state.total_spent = 0.0
if "bulk_drafts" not in st.session_state:
    st.session_state.bulk_drafts = []
if "thread_id" not in st.session_state:
    st.session_state.thread_id = str(uuid.uuid4())

def track_usage(state_values, mock_active):
    if not mock_active and state_values and "usage" in state_values:
        usage = state_values.get("usage", {"input": 0, "output": 0})
        cost = (usage['input'] * 0.000003) + (usage['output'] * 0.000015)
        st.session_state.total_spent += cost

def get_config():
    return {"configurable": {"thread_id": st.session_state.thread_id}}

st.title("✉️ MailMind: Professional Email Agent")

# --- SIDEBAR ---
with st.sidebar:
    st.header("💳 Usage Monitor")
    is_mock = st.toggle("Free Testing Mode (MOCK)", value=True)
    st.metric("Spent", f"${st.session_state.total_spent:.4f}")
    
    st.divider()
    st.header("⚙️ Settings & Mail Test")
    sender_email = os.getenv("SENDER_EMAIL", "Not Configured")
    st.write(f"**Sender:** `{sender_email}`")
    
    if st.button("Verify Mail Credentials 🔗"):
        success, message = test_smtp_connection()
        if success: st.success(message)
        else: st.error(message)

    st.divider()
    st.header("📂 Bulk Processing")
    sample_csv = "recipient_email,recipient_name,subject,goal,tone\ntest@example.com,John,Meeting Request,Schedule a demo,Formal"
    st.download_button("Download Template CSV 📥", data=sample_csv, file_name="mailmind_template.csv")
    uploaded_file = st.file_uploader("Upload filled Template", type=["csv"])

# --- MAIN INTERFACE ---
tab1, tab2 = st.tabs(["Single Draft", "Bulk Processing"])

with tab1:
    col_a, col_b = st.columns(2)
    with col_a:
        tone = st.selectbox("Select Tone", ["Formal", "Friendly", "Assertive", "Urgent"], key="single_tone")
        recipient = st.text_input("Recipient Name", placeholder="e.g. John Doe", key="single_name")
    with col_b:
        email_to = st.text_input("Recipient Email (for Direct Send)", key="single_email")
        mail_subject = st.text_input("Email Subject", placeholder="e.g. Project Update", key="single_subject")
        
    prompt = st.text_area("Email Goal (Describe what you want to say)", height=150, key="single_prompt")

    if st.button("Generate ✨", type="primary"):
        if prompt and recipient:
            st.session_state.thread_id = str(uuid.uuid4())
            initial_state = {"prompt": prompt, "tone": tone, "recipient": recipient, "is_mock": is_mock}
            final_state = mailmind_app.invoke(initial_state, get_config())
            track_usage(final_state, is_mock)
            st.rerun()
        else:
            st.warning("Please fill in the Recipient Name and Goal.")

    current_state = mailmind_app.get_state(get_config())
    
    if current_state.next:
        draft = current_state.values.get("draft", "")
        st.subheader("📝 Review & Edit Draft")
        edited = st.text_area("Final Polish:", value=draft, height=300, key="edit_box")
        
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("✅ Approve & Finalize"):
                mailmind_app.update_state(get_config(), {"draft": edited})
                mailmind_app.invoke(None, get_config())
                st.rerun()
        with c2:
            if st.button("🚀 Approve & Send Now"):
                if not email_to or not mail_subject:
                    st.error("Please enter Email and Subject first!")
                else:
                    mailmind_app.update_state(get_config(), {"draft": edited})
                    mailmind_app.invoke(None, get_config())
                    if is_mock:
                        st.warning("MOCK MODE: Not sent.")
                    else:
                        success, msg = send_real_email(email_to, mail_subject, edited)
                        if success: st.success("Email Sent!")
                        else: st.error(f"Error: {msg}")
        with c3:
            if st.button("🔄 Regenerate"):
                mailmind_app.invoke({"prompt": prompt, "tone": tone, "recipient": recipient, "is_mock": is_mock}, get_config())
                st.rerun()
            
    elif "draft" in current_state.values and current_state.values.get("prompt"):
        st.subheader("🚀 Final Validated Draft")
        final_email = current_state.values["draft"]
        st.info(final_email)
        
        d_col1, d_col2, d_col3 = st.columns(3)
        with d_col1:
            doc = Document()
            doc.add_heading('MailMind Draft', 0)
            doc.add_paragraph(f"To: {recipient}")
            doc.add_paragraph(f"Subject: {mail_subject}")
            doc.add_paragraph(final_email)
            bio_word = io.BytesIO()
            doc.save(bio_word)
            st.download_button("📥 Download Word", data=bio_word.getvalue(), file_name="Email_Draft.docx")
            
        with d_col2:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            safe_text = final_email.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, safe_text)
            st.download_button("📥 Download PDF", data=pdf.output(dest='S').encode('latin-1'), file_name="Email_Draft.pdf")
            
        with d_col3:
            if st.button("🆕 Start New Email"):
                st.session_state.thread_id = str(uuid.uuid4())
                st.rerun()

with tab2:
    if uploaded_file:
        df = pd.read_csv(uploaded_file)
        df.columns = df.columns.str.strip().str.lower()
        required = ['recipient_email', 'recipient_name', 'subject', 'goal', 'tone']
        
        if not all(c in df.columns for c in required):
            st.error(f"CSV missing columns. Required: {', '.join(required)}")
        else:
            if st.button("Step 1: Generate All Drafts 🤖"):
                st.session_state.bulk_drafts = [] 
                progress = st.progress(0)
                for i, row in df.iterrows():
                    res = mailmind_app.invoke(
                        {"prompt": row['goal'], "tone": row['tone'], "recipient": row['recipient_name'], "is_mock": is_mock},
                        {"configurable": {"thread_id": f"bulk-{uuid.uuid4()}"}}
                    )
                    st.session_state.bulk_drafts.append({
                        "email": row['recipient_email'], 
                        "name": row['recipient_name'], 
                        "subject": row['subject'],
                        "draft": res.get("draft", "Error generating draft")
                    })
                    progress.progress((i + 1) / len(df))
                st.success(f"Generated {len(st.session_state.bulk_drafts)} drafts!")

        if st.session_state.bulk_drafts:
            st.divider()
            st.subheader("📝 Review & Edit Bulk Drafts")
            
            # Use a loop to display and allow editing
            for idx, item in enumerate(st.session_state.bulk_drafts):
                with st.expander(f"Review: {item['name']} - {item['subject']}"):
                    # These updates sync directly to session_state
                    st.session_state.bulk_drafts[idx]['subject'] = st.text_input(
                        "Subject:", value=item['subject'], key=f"subj_{idx}"
                    )
                    st.session_state.bulk_drafts[idx]['draft'] = st.text_area(
                        "Body:", value=item['draft'], key=f"body_{idx}", height=200
                    )

            col_b1, col_b2, col_b3 = st.columns(3)
            with col_b1:
                if st.button("🚀 Send All Now"):
                    if is_mock: 
                        st.warning("MOCK MODE: No mail sent.")
                    else:
                        for item in st.session_state.bulk_drafts:
                            send_real_email(item['email'], item['subject'], item['draft'])
                        st.success("All emails dispatched!")
            
            with col_b2:
                # Word Export - Uses the edited session_state
                doc = Document()
                doc.add_heading('MailMind Bulk Export', 0)
                for item in st.session_state.bulk_drafts:
                    doc.add_heading(f"To: {item['name']} ({item['email']})", level=1)
                    doc.add_paragraph(f"Subject: {item['subject']}")
                    doc.add_paragraph(item['draft'])
                    doc.add_page_break()
                bio_word = io.BytesIO()
                doc.save(bio_word)
                st.download_button("📥 Download Word (.docx)", data=bio_word.getvalue(), file_name="Bulk_Emails.docx", key="dl_word_tab2")

            with col_b3:
                # PDF Export - Uses the edited session_state
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                for item in st.session_state.bulk_drafts:
                    pdf.add_page()
                    pdf.set_font("Arial", 'B', 14)
                    pdf.cell(0, 10, f"To: {item['name']} ({item['email']})", ln=True)
                    pdf.set_font("Arial", 'I', 12)
                    pdf.cell(0, 10, f"Subject: {item['subject']}", ln=True)
                    pdf.ln(5)
                    pdf.set_font("Arial", size=12)
                    safe_text = item['draft'].encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 10, safe_text)
                
                pdf_bytes = pdf.output(dest='S').encode('latin-1')
                st.download_button("📥 Download PDF (.pdf)", data=pdf_bytes, file_name="Bulk_Emails.pdf", key="dl_pdf_tab2")

